#!/usr/local/bin/python3

import docx
import re

import indig_parl_utils as utils

from indig_parl_logger import get_logger

docs_logger = get_logger("Process_DOCX_Hansards",
                         a_log_file='logs/proc_docs_hansards_debug.log')


def get_oral_q_doc_obj(docx_path):
    docx_obj = docx.Document(docx_path)
    oral_q = None
    find_start = False
    copy_para = False
    for para in docx_obj.paragraphs:
        if para.style.name == 'Heading 1' and 'oral questions' in para.text.lower():
            docs_logger.debug('Oral Questions header found: %s' % para.text)
            oral_q = docx.Document()
            oral_q.add_paragraph(para.text, style=para.style)
            find_start = True
            copy_para = True
            continue
        if copy_para == True and not para.style.name == 'Heading 1':
            oral_q.add_paragraph(para.text, style=para.style)
        elif para.style.name == 'Heading 1' and find_start:
            break

    return oral_q


def get_oral_q_dialogs(docx_obj):
    dialog_doc_obj_list = []
    copy_para = False
    dialog = docx.Document()
    for para in docx_obj.paragraphs:
        if para.style.name == 'Heading 2':
            if copy_para == True:
                dialog_doc_obj_list.append(dialog)
                dialog = docx.Document()
            else:
                copy_para = True
            dialog.add_paragraph(para.text, style=para.style)
        elif copy_para:
            dialog.add_paragraph(para.text, style=para.style)
    dialog_doc_obj_list.append(dialog)
    docs_logger.debug('Number of oral questions: %s' %
                      len(dialog_doc_obj_list))
    return dialog_doc_obj_list


def get_doc_obj_txt(doc_obj):
    full_text = []
    for para in doc_obj.paragraphs:
        full_text.append(para.text)
    return full_text


def process_doc_oral_q(docx_obj, path, session):
    oral_q_dialogs = get_oral_q_dialogs(docx_obj)
    [just_filename] = path.split('/')[-1:]
    [just_filetitle] = just_filename.split('.')[:-1]
    dialog_txt_file = session+'-'+just_filetitle+'.txt'
    csv_filename = session+'-'+just_filetitle+'.csv'
    doc_text = []
    for dialog in oral_q_dialogs:
        doc_text += get_doc_obj_txt(dialog)
    utils.send_text_to_file('tmp/'+dialog_txt_file,
                            doc_text, data_type='list')
    docs_logger.debug('Oral questions raw text saved to %s' %
                      'tmp/'+dialog_txt_file)

    csv_list = []
    prev_question = False
    q_col = ''
    speaker = ''
    dialogue = ''
    speaker_pattern = r'((?:M[SR]S{0,1}\.|HON\.|HONOURABLE).*?):'
    for dialog in oral_q_dialogs:
        for para in dialog.paragraphs:
            docs_logger.debug(para.text)
            if para.style.name == 'Heading 2':
                docs_logger.debug('New oral question.')
                if prev_question:
                    csv_list.append((q_col, speaker, dialogue))
                    docs_logger.debug('Previous line stored')
                q_col = para.text.replace('\n', ' ')
            else:
                match = re.compile(speaker_pattern).search(para.text)
                if match:
                    if prev_question:
                        csv_list.append((q_col, speaker, dialogue))
                        docs_logger.debug('Previous question.')
                    else:
                        prev_question = True
                    _, speaker, dialogue = re.split(speaker_pattern, para.text,
                                                    maxsplit=1)
                    docs_logger.debug(
                        'Speaker:[%s] Dialog chars: [%s] prev_question:[%s]' %
                        (speaker, len(dialogue), prev_question))
                else:
                    dialogue += ' ' + para.text
    utils.csv_from_list('csvs/'+csv_filename, csv_list,
                        header_row=['Question', 'Speaker', 'Speech'])
