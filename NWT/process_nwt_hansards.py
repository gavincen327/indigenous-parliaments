#!/usr/local/bin/python3

import csv
import docx
import requests
import os
import re
import logging
import sys

import pandas as pd

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import BytesIO


def get_file_list(directory, ext='all-files'):
    files = os.listdir(directory)
    for file in files:
        if not os.path.isfile(file):
            files.pop(files.index(file))
    if not ext == 'all-files':
        for file in files:
            if not file.endswith(ext):
                files.pop(files.index(file))
    return files


def get_source_links(source_file):
    location_dict = {}
    lines = 0
    count = 0
    with open(source_file, newline='') as csv_f:
        csv_reader = csv.reader(csv_f, delimiter=',', quotechar='"')
        for row in csv_reader:
            if lines == 0:
                lines += 1
            else:
                word = row[3]
                pdf = row[4]
                session = row[0] + '-' + row[1]
                location_dict['['+str(count)+']'+row[2]] = session, word, pdf
                count += 1

    return location_dict


def get_pdf_text(web_loc, date, directory="NWT/pdfs/"):
    """
    date should be in following format: 180601
    """
    [name] = web_loc.split('/')[-1:]
    r = requests.get(web_loc)
    file_loc = directory+"["+date+"]"+name
    with open(file_loc, 'wb') as f:
        f.write(r.content)
    return file_loc


def get_docx(web_loc, date, directory="NWT/docs/"):
    """
    date should be in following format: 180601
    """
    [name] = web_loc.split('/')[-1:]
    r = requests.get(web_loc)
    file_loc = directory+"["+date+"]"+name
    with open(file_loc, 'wb') as f:
        f.write(r.content)
    return file_loc


def get_docx_text(file):
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text + '\n')
    return full_text


def get_doc_obj_txt(doc_obj):
    full_text = []
    for para in doc_obj.paragraphs:
        full_text.append(para.text)
    return full_text


def get_oral_q_doc_obj(docx_obj):
    oral_q = None
    find_start = False
    copy_para = False
    for para in docx_obj.paragraphs:
        if para.style.name == 'Heading 1' and 'Oral Questions' == para.text:
            oral_q = docx.Document()
            oral_q.add_paragraph(para.text, style=para.style)
            find_start = True
            copy_para = True
            continue
        if copy_para == True and not para.style.name == 'Heading 1':
            oral_q.add_paragraph(para.text, style=para.style)
        elif para.style.name == 'Heading 1' and find_start:
            break

    return oral_q


def get_oral_q_dialogs(docx_obj):
    dialog_doc_obj_list = []
    copy_para = False
    dialog = docx.Document()
    for para in docx_obj.paragraphs:
        if para.style.name == 'Heading 2':
            if copy_para == True:
                dialog_doc_obj_list.append(dialog)
                dialog = docx.Document()
            else:
                copy_para = True
            dialog.add_paragraph(para.text, style=para.style)
        elif copy_para:
            dialog.add_paragraph(para.text, style=para.style)
    dialog_doc_obj_list.append(dialog)
    logging.debug('Number of oral questions: %s' % len(dialog_doc_obj_list))
    return dialog_doc_obj_list


def send_text_to_file(file_path, data, header='', data_type='text'):
    with open(file_path, 'w') as f:
        if header:
            f.write(header.title() + '\n\n')
        if data_type == 'list':
            for line in data:
                f.write(line + '\n')
        elif data_type == 'dict':
            for item in data.keys():
                f.write(item + ': ' + data[item] + '\n')
        elif data_type == 'tup_list':
            for (a, b) in data:
                f.write(a+' - '+b+'\n')
        else:
            f.write(data)


def download_hansards(links_csv_file):
    # preferred in ['docs','pdfs','both']
    links_filename = links_csv_file.split('.')[0]
    hansards_path_dict = get_source_links(links_csv_file)
    doc_paths = {}
    pdf_paths = {}
    txt_out = []
    print('Downloading Hansards ')
    for key in hansards_path_dict.keys():
        [dte] = key.split(']')[-1:]
        dte_lst = dte.split('-')
        dte_ok = True
        for idx in range(3):
            if len(dte_lst[idx]) < 2:
                dte_lst[idx] = '0' + dte_lst[idx]
                dte_ok = False
        if not dte_ok:
            dte = '-'.join(dte_lst)

        doc_location = ''
        pdf_location = ''
        valid = hansards_path_dict[key][1].endswith('.docx')
        if valid:
            doc_location = get_docx(hansards_path_dict[key][1], dte)
            doc_paths[key] = (hansards_path_dict[key][0],
                              hansards_path_dict[key][1],
                              doc_location)
        elif hansards_path_dict[key][2].endswith('.pdf'):
            pdf_location = get_pdf_text(hansards_path_dict[key][2], dte)
            pdf_paths[key] = (hansards_path_dict[key][0],
                              hansards_path_dict[key][2],
                              pdf_location)
        txt_line = ','.join([hansards_path_dict[key][0],
                             hansards_path_dict[key][1], doc_location,
                             hansards_path_dict[key][2], pdf_location])
        txt_out.append(txt_line)
        print('^')
        logging.debug('\t%s\n\t\t%s:%s\n\t\t%s:%s' %
                      (key, hansards_path_dict[key][1], doc_location,
                       hansards_path_dict[key][2], pdf_location))
    send_text_to_file(links_filename + '_docs.csv', txt_out, data_type='list')
    print('List of downloaded handards created:', links_filename + '_docs.csv')

    return doc_paths, pdf_paths


def process_doc_oral_q(docx_obj, path, session):
    oral_q_dialogs = get_oral_q_dialogs(docx_obj)
    [just_filename] = path.split('/')[-1:]
    [just_filetitle] = just_filename.split('.')[:-1]
    dialog_txt_file = session+'-'+just_filetitle+'.txt'
    csv_filename = session+'-'+just_filetitle+'.csv'
    doc_text = []
    for dialog in oral_q_dialogs:
        doc_text += get_doc_obj_txt(dialog)
    send_text_to_file('NWT/tmp/'+dialog_txt_file, doc_text, data_type='list')
    logging.debug('Oral questions raw text saved to %s' %
                  'NWT/tmp/'+dialog_txt_file)
    csv_list = []

    prev_question = False
    q_col = ''
    speaker = ''
    dialogue = ''
    speaker_pattern = r'((?:M[SR]S{0,1}\.|HON\.|HONOURABLE).*?):'
    for dialog in oral_q_dialogs:
        for para in dialog.paragraphs:
            logging.debug(para.text)
            if para.style.name == 'Heading 2':
                logging.debug('New oral question.')
                if prev_question:
                    csv_list.append((q_col, speaker, dialogue))
                    logging.debug('Previous line stored')
                q_col = para.text.replace('\n', ' ')
            else:
                match = re.compile(speaker_pattern).search(para.text)
                if match:
                    if prev_question:
                        csv_list.append((q_col, speaker, dialogue))
                        logging.debug('Previous question.')
                    else:
                        prev_question = True
                    _, speaker, dialogue = re.split(speaker_pattern, para.text,
                                                    maxsplit=1)
                    logging.debug(
                        'Speaker:[%s] Dialog chars: [%s] prev_question:[%s]' %
                        (speaker, len(dialogue), prev_question))
                else:
                    dialogue += ' ' + para.text
    with open('NWT/csvs/'+csv_filename, mode='w') as csv_file:
        csv_writer = csv.writer(csv_file, delimiter=',', quotechar='"',
                                quoting=csv.QUOTE_ALL)
        csv_writer.writerow(['Question', 'Speaker', 'Speech'])
        for line in csv_list:
            csv_writer.writerow([line[0], line[1], line[2]])


def pdf_to_text(path):
    """Code source: https://towardsdatascience.com/pdf-preprocessing-with-python-19829752af9f

    Arguments:
        path {[type]} -- [description]

    Returns:
        [type] -- [description]
    """
    manager = PDFResourceManager()
    retstr = BytesIO()
    layout = LAParams(all_texts=True)
    device = TextConverter(manager, retstr, laparams=layout)
    filepath = open(path, 'rb')
    interpreter = PDFPageInterpreter(manager, device)
    for page in PDFPage.get_pages(filepath, check_extractable=True):
        interpreter.process_page(page)
    text = retstr.getvalue()
    filepath.close()
    device.close()
    retstr.close()
    return text


def prepare_raw_text(text, store_path=''):
    re_pg_header = r'\s[A-Z][a-z]+,{0,1}\s[A-Z][a-z]+\s\d{1,2},\s[1,2][0,9]\d{2}\s+Nunavut Hansard'
    re_pg_footer = r'\n(?:Page){0,1}\s+(?:\d{2,4})\s+\n'

    cleaned_text = re.sub(re_pg_header, ' ', text)
    cleaned_text = re.sub(re_pg_footer, '\n', cleaned_text)
    if store_path:
        send_text_to_file(store_path, cleaned_text)
    cleaned_text = cleaned_text.replace('\n', ' ')

    return cleaned_text


def get_pattern_text(pattern, search_text):
    pat_complied = re.compile(pattern)
    pat_match = pat_complied.search(search_text)
    if pat_match:
        return pat_match.group(1)
    return None


def get_oral_q_df(text, d_str):
    oral_q_sec = r'Item 6: Oral Questions\s+(.*?)(?:\s+Item \d{1,2}:)'
    q_titles = r'\s*(Question\s+\d{1,3}\s*\S\s\d\(\d\)):'
    section_title = r'Item \d{1,2}:'
    speaker_headers = r'\s*((?:M[r|s]s{0,1}\.{0,1}|Hon\.{0,1}|Honourable|Speaker)(?:\s+(?:O\S){0,1}[A-Z]\w+?\.{0,1}){0,2})\s*(?:\(\w+\)){0,1}:'

    oral_questions = get_pattern_text(oral_q_sec, text)
    if oral_questions:
        df_list = []
        title = 'Item 6: Oral Questions'
        question_list = re.split(q_titles, text)
        for idx in range(len(question_list)):
            if title in question_list[idx]:
                logging.debug('Title in idx: %s' % idx)
                # print(idx, 'entry:', question_list[idx])
                if idx > 0:
                    logging.debug('Previous entry: %s' % question_list[idx-1])
                if idx < len(question_list)-1:
                    logging.debug('Next entry: %s' % question_list[idx+1])
                if idx == 0:
                    question_list = question_list[1:]
                else:
                    question_list = question_list[idx-1:]
                    question_list[1] = question_list[idx].replace(title, '')
                break
        last_item = question_list[len(question_list)-1]
        split_l_item = re.split(section_title, last_item)
        question_list[len(question_list)-1] = split_l_item[0]
        send_text_to_file('NWT/tmp/'+d_str + 'pre_q_split_raw.txt',
                          question_list, data_type='list')
        for idx in range(len(question_list)):
            if not 'Question' == question_list[idx][:8]:
                speaker_pattern = re.compile(speaker_headers)
                speaker_match = speaker_pattern.search(question_list[idx])
                if speaker_match:
                    (start, _) = speaker_match.span()
                    q_tail = question_list[idx][:start]
                    # print('Question tail:', q_tail)
                    dialog = question_list[idx][start:]
                    # print('Old question:', question_list[idx-1])
                    question_list[idx-1] += ': ' + q_tail
                    # print('New question:', question_list[idx-1])
                    question_list[idx] = dialog
        send_text_to_file('NWT/tmp/'+d_str + 'q_split_raw.txt', question_list,
                          data_type='list')
        even_q_list = question_list[1::2]
        odd_q_list = question_list[::2]
        q_dict = dict(zip(odd_q_list, even_q_list))
        logging.debug('split oral questions by question')
        for q in q_dict.keys():
            diag = q_dict[q]
            diag_list = re.split(speaker_headers, diag)
            if diag_list[0] == '':
                diag_list.pop(0)
            even_diag_list = diag_list[1::2]
            odd_diag_list = diag_list[::2]
            diag_tup_list = list(zip(odd_diag_list, even_diag_list))
            q_col_list = [q] * len(diag_tup_list)
            combined = list(
                zip(q_col_list, odd_diag_list, even_diag_list))
            partial_df = pd.DataFrame(
                combined, columns=['question', 'speaker', 'speech'])
            df_list.append(partial_df)
        logging.debug('split question dialogs and dialog splits')
        return pd.concat(df_list)
    else:
        print('Error processing PDF Oral Questions section.')
        logging.debug('Error processing PDF Oral Questions section.')
        return pd.DataFrame(columns=['question', 'speaker', 'speech'])


def process_pdf(pdf_loc, date_str, coding='utf-8'):
    try:
        raw_text = pdf_to_text(pdf_loc)
        raw_text = raw_text.decode(encoding=coding)
        # send_text_to_file('Nunavut/clean_csvs/'+date_str +
        #                   '-raw_text.txt', raw_text)
        # print('raw text sent')
        store_to = 'NWT/tmp/'+date_str + 'rem_hd_ft_raw_text.txt'
        rm_hd_nl_text = prepare_raw_text(raw_text, store_path=store_to)
        # send_text_to_file('Nunavut/clean_csvs/'+date_str +
        #                   'rem_nl_raw_text.txt', rm_hd_nl_text)
        # print('Raw text no newlines sent')
        if rm_hd_nl_text.find("Item 6: Oral Questions") > -1:
            send_text_to_file('NWT/tmp/'+date_str +
                              '-raw_text.txt', raw_text)
            logging.debug('Raw pdf text sent %s' % date_str)
            send_text_to_file('NWT/tmp/'+date_str + 'rem_nl_raw_text.txt',
                              rm_hd_nl_text)
            logging.debug('Raw text no newlines sent %s' % date_str)
            return get_oral_q_df(rm_hd_nl_text, date_str)
        else:
            print('Oral Questions section not present in PDF %s.' % pdf_loc)
            return pd.DataFrame(columns=['question', 'speaker', 'speech'])
    except Exception as e:
        print('ERROR: Error reading/accessing pdf document for %s. [%s].' %
              (date_str, pdf_loc))
        logging.debug('ERROR: Error accessing pdf document for %s. [%s].' %
                      (date_str, pdf_loc))
        logging.debug(e)
        return pd.DataFrame(columns=['question', 'speaker', 'speech'])


def main():
    logging.basicConfig(filename='NWT/proc_nwt_debug.log', level=logging.DEBUG)

    # hansard_csv = 'NWT/nwt_hansards_small_assem.csv'
    # hansard_csv = 'NWT/nwt_hansards_19_assem.csv'
    # hansard_csv = 'NWT/nwt_hansards_18_assem.csv'
    # hansard_csv = 'NWT/nwt_hansards_17_assem.csv'
    # hansard_csv = 'NWT/nwt_hansards_16_assem.csv'
    # hansard_csv = 'NWT/nwt_hansards_15_assem.csv'
    hansard_csv = 'NWT/nwt_hansards_14_assem.csv'

    docx_paths, pdf_paths = download_hansards(hansard_csv)

    print('Processing Hansards ')
    if len(docx_paths.keys()) > 0:
        logging.debug('Processing %s ".docx" hansards.' %
                      len(docx_paths.keys()))
    for key in docx_paths.keys():
        path = docx_paths[key][2]
        logging.debug('Document path: %s' % path)
        if path:
            doc_obj = docx.Document(path)
            oral_q_doc = get_oral_q_doc_obj(doc_obj)
            if not oral_q_doc == None:
                print('(|)')
                logging.debug('Oral questions present.')
                process_doc_oral_q(oral_q_doc, path, docx_paths[key][0])
            else:
                print('(-)')
        else:
            print('(-)')
    if len(pdf_paths.keys()) > 0:
        logging.debug('Processing %s ".pdf" hansards.' %
                      len(pdf_paths.keys()))
    for key in pdf_paths.keys():
        path = pdf_paths[key][2]
        dte = path.split('/')[-1:][0][1:11]
        if path:
            hansard_df = process_pdf(path, dte)
            if not hansard_df.empty:
                print('(|)')
                hansard_df.to_csv("Nunavut/clean_csvs/"+dte + ".csv",
                                  encoding='utf-8-sig')
            else:
                print('(|)')

    print('Extracted Oral Questions in folder "NWT/csvs/"')


if __name__ == '__main__':
    main()
